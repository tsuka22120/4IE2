#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_all.py

このスクリプトは以下を自動生成する。
 - Wordレポート (report_assignment1.docx) : python-docx を使用
 - Verilog ソース (fa_fs_top.v)
 - Verilog テストベンチ (fa_fs_tb.v)
 - ピン割り当てファイル (gowin_pins.cst)

実行環境:
 - python3
 - python-docx (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os

# ----------------------------
# ファイル名定義
# ----------------------------
REPORT_FILENAME = "report_assignment1.docx"
VERILOG_FILENAME = "fa_fs_top.v"
TB_FILENAME = "fa_fs_tb.v"
CST_FILENAME = "gowin_pins.cst"

# ----------------------------
# Verilog ソース（階層設計: HA->FA, HS->FS, top モジュール）
# ----------------------------
verilog_code = r"""
// fa_fs_top.v
// Gowin GW1NR-9 target
// 入力ピン: SW1(69)=iA, SW2(68)=iB, SW3(57)=iCy/iBr, SW4(56)=mode
// 出力ピン: LD1(10)=sum/diff (表示反転), LD2(11)=carry/borrow (表示反転)
// SW4 = 0 -> 全加算器 (FA)
// SW4 = 1 -> 全減算器 (FS)
// 表示LEDは反転表示されるため、出力を反転している。

`timescale 1ns/1ps

module ha(
    input  wire A,
    input  wire B,
    output wire S,
    output wire C
);
    // Half-adder: S = A xor B, C = A & B
    assign S = A ^ B;
    assign C = A & B;
endmodule

module fa(
    input  wire A,
    input  wire B,
    input  wire Cin,
    output wire S,
    output wire Cout
);
    // FA implemented by two HAs
    wire s1, c1, c2;
    ha ha1(.A(A), .B(B), .S(s1), .C(c1));
    ha ha2(.A(s1), .B(Cin), .S(S), .C(c2));
    assign Cout = c1 | c2;
endmodule

module hs(
    input  wire A,
    input  wire B,
    output wire D,   // difference
    output wire Br   // borrow
);
    // Half-subtractor: D = A xor B, Br = (~A) & B
    assign D  = A ^ B;
    assign Br = (~A) & B;
endmodule

module fs(
    input  wire A,
    input  wire B,
    input  wire Bin,
    output wire D,
    output wire Bout
);
    // Full subtractor implemented by two HS
    wire d1, b1, b2;
    hs hs1(.A(A), .B(B), .D(d1), .Br(b1));
    hs hs2(.A(d1), .B(Bin), .D(D),  .Br(b2)); // note: second HS subtracts Bin from d1
    assign Bout = b1 | b2;
endmodule

module fa_fs_top(
    input  wire iA,    // SW1 (pin 69)
    input  wire iB,    // SW2 (pin 68)
    input  wire iC,    // SW3 (pin 57) : carry-in or borrow-in
    input  wire mode,  // SW4 (pin 56) : 0=FA, 1=FS
    output wire LD1,   // LD1 (pin 10) : sum/diff (display inverted)
    output wire LD2    // LD2 (pin 11) : carry/borrow (display inverted)
);
    wire s_sum, c_cout;
    wire s_diff, b_bout;
    // instantiate FA and FS
    fa fa_inst(.A(iA), .B(iB), .Cin(iC), .S(s_sum), .Cout(c_cout));
    fs fs_inst(.A(iA), .B(iB), .Bin(iC), .D(s_diff), .Bout(b_bout));

    // multiplex by mode: mode==0 -> FA, mode==1 -> FS
    wire out_bit, out_flag;
    assign out_bit  = (mode == 1'b0) ? s_sum : s_diff;
    assign out_flag = (mode == 1'b0) ? c_cout : b_bout;

    // LED displays are inverted => invert signals
    assign LD1 = ~out_bit;
    assign LD2 = ~out_flag;
endmodule
"""

# ----------------------------
# Testbench
# ----------------------------
tb_code = r"""
// fa_fs_tb.v
`timescale 1ns/1ps
module fa_fs_tb;
    reg A, B, C, mode;
    wire LD1, LD2;
    // Instantiate DUT
    fa_fs_top dut(.iA(A), .iB(B), .iC(C), .mode(mode), .LD1(LD1), .LD2(LD2));

    integer i;
    reg [2:0] vec;
    initial begin
        $dumpfile("fa_fs_tb.vcd");
        $dumpvars(0, fa_fs_tb);
        $display("time\t mode A B C | LD1(sum/diff_inv) LD2(carry/borrow_inv) | real_out sum/diff carry/borrow");
        $display("--------------------------------------------------------------------------");
        for (mode = 0; mode <= 1; mode = mode + 1) begin
            for (i = 0; i < 8; i = i + 1) begin
                vec = i;
                A = vec[2];
                B = vec[1];
                C = vec[0];
                #5;
                // LD outputs are inverted on hardware, so real outputs are inverted-back:
                $display("%g\t  %b   %b %b %b |  %b                %b            |  %b",
                         $time, mode, A, B, C, LD1, LD2, 1'b0);
                // For clarity, also print computed expected values:
                if (mode == 0) begin
                    // FA expected
                    $display("  (FA) A=%b B=%b Cin=%b => Sum=%b Cout=%b (LD shows inverted).",
                        A,B,C, (A ^ B) ^ C, ((A & B) | ((A ^ B) & C)));
                end else begin
                    // FS expected: Difference = A ^ B ^ Bin ; Borrow = ((~A)&B) | (((~(A ^ B)) & C))
                    reg diff;
                    reg bout;
                    diff = (A ^ B) ^ C;
                    bout = ((~A) & B) | (((~(A ^ B)) & C));
                    $display("  (FS) A=%b B=%b Bin=%b => Diff=%b Bout=%b (LD shows inverted).",
                        A,B,C, diff, bout);
                end
            end
        end
        #10;
        $finish;
    end
endmodule
"""

# ----------------------------
# CST (pin assignment) - simple set_io style
# ----------------------------
cst_code = """# gowin_pins.cst
# Pin assignment (simple set_io format)
# Device: Gowin GW1NR-9 (example)
# Inputs:
set_io iA 69   # SW1
set_io iB 68   # SW2
set_io iC 57   # SW3 (carry-in / borrow-in)
set_io mode 56 # SW4 (mode: 0=FA,1=FS)
# Outputs:
set_io LD1 10  # LD1: sum/diff (inverted on display)
set_io LD2 11  # LD2: carry/borrow (inverted on display)

# Note:
# - 実際の Gowin EDA で使う制約ファイル形式はバージョンによって異なる。
# - 必要に応じて、プロジェクトの .pcf や .xdc 形式へ変換すること。
"""

# ----------------------------
# レポート本文テンプレート（日本語・だ・である調）
# ----------------------------
def generate_report():
    doc = Document()

    # 日本語フォント設定 (段落・ラン単位でフォントを設定)
    style = doc.styles['Normal']
    style.font.name = 'MS Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ ゴシック')
    style.font.size = Pt(10.5)

    # 表紙（1ページ目）
    doc.add_page_break()  # ensure we can place cover as first page formatting
    # Actually python-docx does not support easy 'cover page' control; we will build first page
    cover = doc.sections[0]
    # Title block:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("課題提出表紙\n")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'MS Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ ゴシック')

    doc.add_paragraph("")  # spacer
    # metadata table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    # left column labels, right column values
    rows = [
        ("提出日", datetime.now().strftime("%Y-%m-%d")),
        ("氏名", "22120 塚田 勇人"),
        ("講義名", "集積回路設計 演習1"),
        ("課題名", "課題1: 全加算器および全減算器の階層設計"),
        ("報告事項", "回路設計、Verilog実装、検証、ピン割り当て"),
        ("備考", "LDは出力と表示が反転するため、Verilogで反転を実装")
    ]
    for i, (k, v) in enumerate(rows):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.text = k
        cell_v.text = v
        # set font
        for cell in (cell_k, cell_v):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'MS Gothic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ ゴシック')
                    run.font.size = Pt(11)

    doc.add_page_break()

    # 目次的に章を作成
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'MS Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'ＭＳ ゴシック')

    # 1. 目的
    add_heading("1. 目的", level=1)
    p = doc.add_paragraph("本実験の目的は、半加算器と半減算器を用いた階層設計により、全加算器 (FA) および全減算器 (FS) を実装し、Gowin GW1NR-9 FPGA 上で動作を確認することである。さらに、"
                          "FPGA ボード上のスイッチを用いて加算・減算を切り替え、LED 表示の反転仕様に対応した論理を実演することにある。")
    p.style = doc.styles['Normal']

    # 2. 原理（理論）
    add_heading("2. 原理（理論）", level=1)
    doc.add_paragraph(
        "半加算器 (HA) は入力 A, B に対して和 S = A ⊕ B、桁上げ C = A・B を出力する回路である。"
        "全加算器 (FA) は二つの HA を用いて実現できる。第一段の HA がまず A, B の部分和と中間桁上げを算出し、"
        "第二段の HA が部分和と Cin を合成して最終和 S を得る。最終桁上げは二段の桁上げの論理和として得られる。"
    )
    doc.add_paragraph(
        "半減算器 (HS) は差 D = A ⊕ B、借用 Br = (¬A)・B を出力する。全減算器 (FS) は二つの HS を用いて実装できる。"
        "FS は A - B - Bin を計算し、差分と借用を出力する。各段の借用は論理和により最終借用となる。"
    )
    doc.add_paragraph(
        "また本 FPGA ボードの LED 表示は出力値と表示が反転する仕様であるため、実機接続時には出力線を反転して LED に接続する必要がある。"
        "本設計では Verilog 内で出力論理を反転している。"
    )

    # 3. 使用器具
    add_heading("3. 使用器具", level=1)
    doc.add_paragraph("使用器具は以下である。")
    doc.add_paragraph("・Gowin EDA（論理合成・配置配線）\n・Gowin GW1NR-9 搭載 FPGA ボード\n・波形観測用に ModelSim / iverilog + gtkwave（検証環境）")

    # 4. 実験方法
    add_heading("4. 実験方法", level=1)
    doc.add_paragraph(
        "1. 階層設計として HA/HS を作成し、それらを用いて FA/FS を構築する。\n"
        "2. トップモジュールでモード入力 (SW4) により加算/減算を切替可能とする。\n"
        "3. 出力は LED の反転仕様に合わせて論理反転を行う。\n"
        "4. 合成・配置配線を行い、FPGA 上で動作確認を行う。"
    )

    # 5. 実験結果（コード、ピン割り当て、テストベンチ、波形の説明）
    add_heading("5. 実験結果", level=1)
    doc.add_paragraph("以下に実装したコード、ピン割り当て、およびテストベンチを示す。コードは 1 行 1 列の表に収めて掲載している。")

    # Insert Verilog code into a 1x1 table
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 1-1)
    # Put the verilog code text in the cell
    cell_paragraph = cell.paragraphs[0]
    # Use a run with monospace font
    run = cell_paragraph.add_run(verilog_code)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)

    # Add pin assignment as a table (caption above for tables)
    # Insert caption (table captions placed above)
    doc.add_paragraph("表 1: ピン割り当て（簡易版）", style='Intense Quote')
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    table2.cell(0,0).text = "信号名"
    table2.cell(0,1).text = "FPGA ピン"
    table2.cell(1,0).text = "iA (SW1)"
    table2.cell(1,1).text = "69"
    table2.cell(2,0).text = "iB (SW2)"
    table2.cell(2,1).text = "68"
    table2.cell(3,0).text = "iC (SW3)"
    table2.cell(3,1).text = "57"
    # next small table for remaining
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    table3.cell(0,0).text = "mode (SW4)"
    table3.cell(0,1).text = "56"
    table3.cell(1,0).text = "LD1 (sum/diff)"
    table3.cell(1,1).text = "10 (表示は反転)"
    table3.cell(2,0).text = "LD2 (carry/borrow)"
    table3.cell(2,1).text = "11 (表示は反転)"

    # Waveform description (text)
    doc.add_paragraph("図 1: 代表的な波形例（説明）", style='Intense Quote')
    doc.add_paragraph(
        "テストベンチにより A,B,C の全組合せを mode=0,1 で順に刺激した。"
        "FA 動作時には S = A ⊕ B ⊕ Cin, Cout = (A & B) | ((A ⊕ B) & Cin) となる。"
        "FS 動作時には D = A ⊕ B ⊕ Bin, Bout = ((~A)&B) | (((~(A ⊕ B))&Bin)) である。"
        "LD はハード的に反転されているため、波形観測時は LD の論理を反転して解釈する必要がある。"
    )

    # 6. 考察
    add_heading("6. 考察", level=1)
    doc.add_paragraph(
        "本実験により、半加算器・半減算器を用いた階層設計により FA/FS を容易に構築できることが確認できた。"
        "特に HA/HS を再利用することで回路の構造が明確になり、検証が容易である。"
        "LED 表示の反転仕様は小さな注意点ではあるが、実機での確認時に見落とすと誤認の原因となるため、"
        "設計段階で反転を明示的に実装しておくことが重要である。"
    )
    doc.add_paragraph(
        "今後の改良点として、さらにビット幅を拡張した複数ビット加減算器への拡張、"
        "および FPGA のタイミング解析を行いクロックドメインでのパイプライン化を検討することが挙げられる。"
    )

    # Save doc
    doc.save(REPORT_FILENAME)
    print(f"Report written to {REPORT_FILENAME}")

# ----------------------------
# ファイル書き出し
# ----------------------------
def write_files():
    # Verilog
    with open(VERILOG_FILENAME, "w", encoding="utf-8") as f:
        f.write(verilog_code.strip() + "\n")
    print(f"Verilog written to {VERILOG_FILENAME}")

    # TB
    with open(TB_FILENAME, "w", encoding="utf-8") as f:
        f.write(tb_code.strip() + "\n")
    print(f"Testbench written to {TB_FILENAME}")

    # CST
    with open(CST_FILENAME, "w", encoding="utf-8") as f:
        f.write(cst_code)
    print(f"CST written to {CST_FILENAME}")

# ----------------------------
# main
# ----------------------------
if __name__ == "__main__":
    generate_report()
    write_files()
    print("All files generated. Please open the Word document and verilog files for review.")
