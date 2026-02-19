"""
集積回路設計演習1 レポート自動生成プログラム
作成者: 塚田 勇人
学籍番号: 22120
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR
import os
from pathlib import Path
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont


class ReportGenerator:
    """レポート自動生成クラス"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
        # 基本情報
        self.author_name = "塚田 勇人"
        self.student_id = "22120"
        self.subject_name = "集積回路設計演習1"
        self.assignment_name = "組合せ回路とHDL記述（全加算器・全減算器の構築）"
        self.output_filename = "Report_22120_Tsukada.docx"
        
        # パス設定
        self.base_dir = Path(__file__).parent
        self.src_dir = self.base_dir / "src"
        self.img_dir = self.base_dir / "img"
        
    def setup_styles(self):
        """スタイルの設定"""
        # 本文スタイル
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'MS 明朝'
        font.size = Pt(10.5)
        
        # 英数字フォント設定
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'MS 明朝')
        style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        
    def add_header_info(self):
        """ヘッダー情報の追加"""
        # 科目名
        p = self.doc.add_paragraph()
        run = p.add_run(f"科目名: {self.subject_name}")
        run.font.size = Pt(12)
        run.font.bold = True
        
        # 課題名
        p = self.doc.add_paragraph()
        run = p.add_run(f"課題名: {self.assignment_name}")
        run.font.size = Pt(12)
        
        # 学籍番号
        p = self.doc.add_paragraph()
        run = p.add_run(f"学籍番号: {self.student_id}")
        run.font.size = Pt(11)
        
        # 氏名
        p = self.doc.add_paragraph()
        run = p.add_run(f"氏名: {self.author_name}")
        run.font.size = Pt(11)
        
        # 空行
        self.doc.add_paragraph()
        
    def add_chapter1_purpose(self):
        """第1章 目的"""
        self.add_heading("第1章 目的", level=1)
        
        content = (
            "本実験の目的は、Gowin社製FPGA GW1NR-9を使用し、"
            "Verilog HDLによる組合せ回路の設計と実装を通じて、"
            "ハードウェア記述言語の基礎を習得することである。"
            "具体的には、半加算器を2つ組み合わせた全加算器と、"
            "半減算器を2つ組み合わせた全減算器を構築する。"
            "さらに、SW4により全加算器と全減算器を切り替える機能を実装し、"
            "FPGAボード上のスイッチとLEDを用いた入出力動作の確認を行う。"
            "入力はSW1(iA)、SW2(iB)、SW3(iCy/iBr)の3つのスイッチで行い、"
            "出力はLD1(Sum/Diff)とLD2(Cout/Br_out)の2つのLEDで確認する。"
            "なお、LEDは出力値と点灯状態が反転することに注意が必要である。"
        )
        
        self.add_paragraph(content)
        
    def add_chapter2_principle(self):
        """第2章 原理・設計"""
        self.add_heading("第2章 原理・設計", level=1)
        
        # 2.1 ブロック図
        self.add_heading("2.1 ブロック図", level=2)
        self.add_paragraph(
            "本回路の全体構成を図1に示す。"
            "全加算器と全減算器は並列に配置され、SW4の入力により出力が切り替えられる。"
        )
        
        # ブロック図の作成と挿入
        block_diagram_path = self.create_block_diagram()
        if block_diagram_path:
            try:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(block_diagram_path), width=Inches(5.5))
                
                # キャプション
                p_caption = self.doc.add_paragraph()
                run = p_caption.add_run("図1: 全加算器・全減算器のブロック図")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = self.doc.add_paragraph("[ブロック図の作成に失敗しました]")
                for run in p.runs:
                    run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # 空行
        
        # 2.2 全加算器の原理
        self.add_heading("2.2 全加算器の原理", level=2)
        self.add_paragraph(
            "全加算器(Full Adder)は、2つの入力ビットiA、iBと桁上がり入力iCyを加算し、"
            "和oSと桁上がり出力oCyを生成する回路である。"
            "半加算器を2段接続することで実現できる。"
            "表1に全加算器の真理値表を示す。"
        )
        
        # 全加算器の真理値表
        self.add_truth_table_full_adder()
        
        # 2.3 全減算器の原理
        self.add_heading("2.3 全減算器の原理", level=2)
        self.add_paragraph(
            "全減算器(Full Subtractor)は、被減数iAから減数iBと借入れiBrを減算し、"
            "差oDと借り出しoBrを生成する回路である。"
            "半減算器を2段接続することで実現できる。"
            "表2に全減算器の真理値表を示す。"
        )
        
        # 全減算器の真理値表
        self.add_truth_table_full_subtractor()
        
    def add_chapter3_implementation(self):
        """第3章 実装"""
        self.add_heading("第3章 実装", level=1)
        
        # 3.1 Verilog HDLコード
        self.add_heading("3.1 Verilog HDLコード", level=2)
        self.add_paragraph(
            "本回路のVerilog HDL実装をリスト1に示す。"
            "半加算器、全加算器、半減算器、全減算器、およびトップモジュールで構成される。"
        )
        
        # main.vの読み込み
        self.add_code_listing(
            "main.v",
            "リスト1: 全加算器・全減算器の実装コード (main.v)"
        )
        
        # 3.2 ピン割り当て
        self.add_heading("3.2 ピン割り当て", level=2)
        self.add_paragraph(
            "FPGA GW1NR-9とデバイス（スイッチ、LED）との接続を表3に示す。"
            "ピン割り当ての詳細はリスト2に示す制約ファイルで定義される。"
        )
        
        # ピン割り当て表
        self.add_pin_assignment_table()
        
        # .cstファイルの読み込み
        self.add_code_listing(
            "test.cst",
            "リスト2: ピン制約ファイル (test.cst)"
        )
        
    def add_chapter4_simulation(self):
        """第4章 シミュレーション結果"""
        self.add_heading("第4章 シミュレーション結果", level=1)
        
        # 4.1 テストベンチ
        self.add_heading("4.1 テストベンチ", level=2)
        self.add_paragraph(
            "動作検証のため、Verilogテストベンチを作成した。"
            "テストベンチでは、SW4=0の場合に全加算器の動作を、"
            "SW4=1の場合に全減算器の動作を確認する。"
            "各入力パターンに対して10ns間隔で信号を変化させ、"
            "出力LED1、LD2の状態を観測した。"
            "テストベンチのコードをリスト3に示す。"
        )
        
        # テストベンチコードの読み込み
        self.add_code_listing(
            "simu.v",
            "リスト3: テストベンチコード (simu.v)"
        )
        
        # 4.2 タイミングチャート
        self.add_heading("4.2 タイミングチャート", level=2)
        self.add_paragraph(
            "シミュレーション結果のタイミングチャートを図2に示す。"
            "全加算器モード(SW4=0)では、入力の加算結果が正しく出力されることを確認した。"
            "全減算器モード(SW4=1)では、減算および借入れ処理が正常に動作することを確認した。"
        )
        
        # タイミングチャートの挿入試行
        self.add_image_or_placeholder(
            "simu.png",
            "図2: シミュレーションのタイミングチャート"
        )
        
    def add_chapter5_discussion(self):
        """第5章 考察"""
        self.add_heading("第5章 考察", level=1)
        
        content = (
            "本実験では、Verilog HDLを用いて全加算器と全減算器を設計・実装し、"
            "シミュレーションにより動作を検証した。\n\n"
            "シミュレーション結果から、SW4=0の場合は全加算器として、"
            "SW4=1の場合は全減算器として正常に動作することが確認できた。"
            "全加算器では、3つの入力の加算と桁上がりが正しく計算され、"
            "全減算器では減算と借入れの処理が適切に行われた。\n\n"
            "また、LEDの出力が反転することを考慮し、"
            "回路設計において出力信号を反転させる処理を実装した。"
            "これにより、論理値と実際のLED点灯状態の対応が正しく保たれる。\n\n"
            "半加算器と半減算器をモジュール化し、"
            "それらを組み合わせて全加算器・全減算器を構成する設計手法により、"
            "コードの可読性と保守性が向上した。"
            "階層的なモジュール設計は、大規模回路の開発において重要である。\n\n"
            "今回の実験を通じて、Verilog HDLによるハードウェア記述の基礎と、"
            "FPGAを用いた組合せ回路の実装方法を習得することができた。"
        )
        
        # 段落ごとに分割して追加
        for para in content.split('\n\n'):
            if para.strip():
                self.add_paragraph(para.strip())
        
    def add_heading(self, text, level=1):
        """見出しの追加"""
        heading = self.doc.add_heading(text, level=level)
        # 見出しのフォント設定
        for run in heading.runs:
            run.font.name = 'MS 明朝'
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
            if level == 1:
                run.font.size = Pt(14)
            elif level == 2:
                run.font.size = Pt(12)
        return heading
        
    def add_paragraph(self, text):
        """段落の追加"""
        p = self.doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'MS 明朝'
            run.font.size = Pt(10.5)
        return p
        
    def add_truth_table_full_adder(self):
        """全加算器の真理値表を追加"""
        # 表のキャプション
        p = self.doc.add_paragraph()
        run = p.add_run("表1: 全加算器の真理値表")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 表の作成
        table = self.doc.add_table(rows=9, cols=5)
        table.style = 'Table Grid'
        
        # ヘッダー行
        headers = ['iA', 'iB', 'iCy', 'oS', 'oCy']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.format_table_cell(cell, bold=True)
        
        # データ行
        data = [
            ['0', '0', '0', '0', '0'],
            ['0', '0', '1', '1', '0'],
            ['0', '1', '0', '1', '0'],
            ['0', '1', '1', '0', '1'],
            ['1', '0', '0', '1', '0'],
            ['1', '0', '1', '0', '1'],
            ['1', '1', '0', '0', '1'],
            ['1', '1', '1', '1', '1'],
        ]
        
        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                self.format_table_cell(cell)
        
        self.doc.add_paragraph()  # 空行
        
    def add_truth_table_full_subtractor(self):
        """全減算器の真理値表を追加"""
        # 表のキャプション
        p = self.doc.add_paragraph()
        run = p.add_run("表2: 全減算器の真理値表")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 表の作成
        table = self.doc.add_table(rows=9, cols=5)
        table.style = 'Table Grid'
        
        # ヘッダー行
        headers = ['iA', 'iB', 'iBr', 'oD', 'oBr']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.format_table_cell(cell, bold=True)
        
        # データ行
        data = [
            ['0', '0', '0', '0', '0'],
            ['0', '0', '1', '1', '1'],
            ['0', '1', '0', '1', '1'],
            ['0', '1', '1', '0', '1'],
            ['1', '0', '0', '1', '0'],
            ['1', '0', '1', '0', '0'],
            ['1', '1', '0', '0', '0'],
            ['1', '1', '1', '1', '1'],
        ]
        
        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                self.format_table_cell(cell)
        
        self.doc.add_paragraph()  # 空行
        
    def add_pin_assignment_table(self):
        """ピン割り当て表を追加"""
        # 表のキャプション
        p = self.doc.add_paragraph()
        run = p.add_run("表3: デバイスとFPGAとの接続")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 表の作成
        table = self.doc.add_table(rows=7, cols=3)
        table.style = 'Table Grid'
        
        # ヘッダー行
        headers = ['デバイス', '信号名', 'ピン番号']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self.format_table_cell(cell, bold=True)
        
        # データ行
        data = [
            ['SW1', 'iA', '69'],
            ['SW2', 'iB', '68'],
            ['SW3', 'iCy/iBr', '57'],
            ['SW4', 'Mode Select', '56'],
            ['LD1', 'Sum/Diff', '10'],
            ['LD2', 'Cout/Br_out', '11'],
        ]
        
        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                self.format_table_cell(cell)
        
        self.doc.add_paragraph()  # 空行
        
    def format_table_cell(self, cell, bold=False):
        """表のセルのフォーマット設定"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'MS 明朝'
                run.font.size = Pt(10)
                run.font.bold = bold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    def add_code_listing(self, filename, caption):
        """コードリストの追加"""
        # キャプション
        p = self.doc.add_paragraph()
        run = p.add_run(caption)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # センタリング
        
        # ファイル読み込み試行
        file_path = self.src_dir / filename
        if file_path.exists():
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    code = f.read()
                
                # コードを枠線付き表に入れる
                table = self.doc.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                cell = table.rows[0].cells[0]
                
                # コードを等幅フォントで設定
                p_code = cell.paragraphs[0]
                run = p_code.add_run(code)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
            except Exception as e:
                # エラー時のメッセージ
                p = self.doc.add_paragraph(f"[ファイルが見つかりません: {filename}]")
                for run in p.runs:
                    run.font.size = Pt(10)
        else:
            # ファイルが存在しない場合
            p = self.doc.add_paragraph(f"[ファイルが見つかりません: {filename}]")
            for run in p.runs:
                run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # 空行
        
    def add_image_or_placeholder(self, filename, caption):
        """画像またはプレースホルダーの追加"""
        file_path = self.img_dir / filename
        
        if file_path.exists():
            try:
                # 画像を追加
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(file_path), width=Inches(5.5))
                
                # キャプション
                p_caption = self.doc.add_paragraph()
                run = p_caption.add_run(caption)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黒色に設定
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                # 画像読み込みエラー
                p = self.doc.add_paragraph(f"[ファイルが見つかりません: {filename}]")
                for run in p.runs:
                    run.font.size = Pt(10)
        else:
            # ファイルが存在しない場合
            p = self.doc.add_paragraph(f"[ファイルが見つかりません: {filename}]")
            for run in p.runs:
                run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # 空行
    
    def create_block_diagram(self):
        """ブロック図を作成（添付画像を参考にした設計）"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # 画像サイズ
            width = 1600
            height = 500
            
            # 画像とドロー用オブジェクト作成
            img = Image.new('RGB', (width, height), 'white')
            draw = ImageDraw.Draw(img)
            
            # フォント設定
            try:
                font_large = ImageFont.truetype("arial.ttf", 20)
                font_medium = ImageFont.truetype("arial.ttf", 16)
                font_small = ImageFont.truetype("arial.ttf", 14)
            except:
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
                font_small = ImageFont.load_default()
            
            # 色設定
            black = (0, 0, 0)
            red = (255, 100, 100)
            green = (100, 200, 100)
            light_gray = (240, 240, 240)
            
            # 座標定義
            # 入力スイッチ（左側）
            sw_start_x = 50
            sw4_y = 60
            sw1_y = 150
            sw2_y = 240
            sw3_y = 330
            sw_width = 60
            sw_height = 40
            
            # Full Adder
            fa_x = 300
            fa_y = 100
            fa_width = 120
            fa_height = 120
            
            # Full Subtractor
            fs_x = 300
            fs_y = 260
            fs_width = 120
            fs_height = 120
            
            # 中間信号ボックス（2ビット幅の信号）
            add_box_x = 560
            add_box_y = 135
            sub_box_x = 560
            sub_box_y = 295
            signal_box_w = 50
            signal_box_h = 60
            
            # マルチプレクサ（三角形）
            mux_x = 800
            mux_y = 180
            mux_width = 80
            mux_height = 160
            
            # NOT gate
            not_x = 1050
            not_y_top = 210
            not_y_bot = 270
            not_size = 40
            
            # 出力LED
            led_x = 1350
            ld1_y = 180
            ld2_y = 300
            led_size = 50
            
            # === 入力スイッチの描画 ===
            switches = [
                (sw_start_x, sw4_y, 'SW4'),
                (sw_start_x, sw1_y, 'SW1'),
                (sw_start_x, sw2_y, 'SW2'),
                (sw_start_x, sw3_y, 'SW3')
            ]
            
            for x, y, label in switches:
                draw.rectangle([x, y, x+sw_width, y+sw_height], outline=black, fill=light_gray, width=2)
                draw.text((x+8, y+12), label, fill=black, font=font_medium)
            
            # === Full Adder の描画 ===
            draw.rectangle([fa_x, fa_y, fa_x+fa_width, fa_y+fa_height], outline=black, fill=(230, 240, 255), width=2)
            draw.text((fa_x+10, fa_y+10), 'fullAdder', fill=black, font=font_medium)
            # 入力ラベル
            draw.text((fa_x-25, fa_y+20), 'iA', fill=black, font=font_small)
            draw.text((fa_x-25, fa_y+50), 'iB', fill=black, font=font_small)
            draw.text((fa_x-30, fa_y+80), 'iCy', fill=black, font=font_small)
            # 出力ラベル
            draw.text((fa_x+fa_width+5, fa_y+40), 'oS', fill=black, font=font_small)
            draw.text((fa_x+fa_width+5, fa_y+70), 'oCy', fill=black, font=font_small)
            # 内部構造（U_FA）
            draw.text((fa_x+30, fa_y+90), 'U_FA', fill=black, font=font_small)
            
            # === Full Subtractor の描画 ===
            draw.rectangle([fs_x, fs_y, fs_x+fs_width, fs_y+fs_height], outline=black, fill=(255, 240, 230), width=2)
            draw.text((fs_x+5, fs_y+10), 'fullSubtractor', fill=black, font=font_small)
            # 入力ラベル
            draw.text((fs_x-25, fs_y+20), 'iA', fill=black, font=font_small)
            draw.text((fs_x-25, fs_y+50), 'iB', fill=black, font=font_small)
            draw.text((fs_x-30, fs_y+80), 'iBr', fill=black, font=font_small)
            # 出力ラベル
            draw.text((fs_x+fs_width+5, fs_y+40), 'oD', fill=black, font=font_small)
            draw.text((fs_x+fs_width+5, fs_y+70), 'oBr', fill=black, font=font_small)
            # 内部構造（U_FS）
            draw.text((fs_x+30, fs_y+90), 'U_FS', fill=black, font=font_small)
            
            # === 配線: SW1, SW2, SW3 から FA と FS へ ===
            # SW1 (分岐点)
            sw1_out = sw_start_x + sw_width
            branch1_x = 200
            draw.line([(sw1_out, sw1_y+20), (branch1_x, sw1_y+20)], fill=red, width=2)
            draw.ellipse([branch1_x-4, sw1_y+16, branch1_x+4, sw1_y+24], fill=red)
            # SW1 -> FA.iA
            draw.line([(branch1_x, sw1_y+20), (branch1_x, fa_y+30), (fa_x, fa_y+30)], fill=red, width=2)
            draw.text((branch1_x+5, sw1_y+5), 'SW1', fill=red, font=font_small)
            # SW1 -> FS.iA
            draw.line([(branch1_x, sw1_y+20), (branch1_x, fs_y+30), (fs_x, fs_y+30)], fill=red, width=2)
            
            # SW2 (分岐点)
            branch2_x = 220
            draw.line([(sw1_out, sw2_y+20), (branch2_x, sw2_y+20)], fill=red, width=2)
            draw.ellipse([branch2_x-4, sw2_y+16, branch2_x+4, sw2_y+24], fill=red)
            # SW2 -> FA.iB
            draw.line([(branch2_x, sw2_y+20), (branch2_x, fa_y+60), (fa_x, fa_y+60)], fill=red, width=2)
            draw.text((branch2_x+5, sw2_y+5), 'SW2', fill=red, font=font_small)
            # SW2 -> FS.iB
            draw.line([(branch2_x, sw2_y+20), (branch2_x, fs_y+60), (fs_x, fs_y+60)], fill=red, width=2)
            
            # SW3 (分岐点)
            branch3_x = 240
            draw.line([(sw1_out, sw3_y+20), (branch3_x, sw3_y+20)], fill=red, width=2)
            draw.ellipse([branch3_x-4, sw3_y+16, branch3_x+4, sw3_y+24], fill=red)
            # SW3 -> FA.iCy
            draw.line([(branch3_x, sw3_y+20), (branch3_x, fa_y+90), (fa_x, fa_y+90)], fill=red, width=2)
            draw.text((branch3_x+5, sw3_y+5), 'SW3', fill=red, font=font_small)
            # SW3 -> FS.iBr
            draw.line([(branch3_x, sw3_y+20), (branch3_x, fs_y+90), (fs_x, fs_y+90)], fill=red, width=2)
            
            # === 中間信号ボックス (addResult, subResult) ===
            # addResult [1:0]
            draw.rectangle([add_box_x, add_box_y, add_box_x+signal_box_w, add_box_y+signal_box_h], 
                          outline=black, fill=light_gray, width=2)
            draw.text((add_box_x+8, add_box_y+5), '0', fill=black, font=font_medium)
            draw.text((add_box_x+8, add_box_y+30), '1', fill=black, font=font_medium)
            draw.text((add_box_x-80, add_box_y+20), 'addResult', fill=black, font=font_small)
            
            # FA出力 -> addResult
            # oS -> bit 0
            draw.line([(fa_x+fa_width, fa_y+50), (add_box_x, add_box_y+15)], fill=red, width=2)
            draw.text((fa_x+fa_width+40, fa_y+35), 'addC', fill=red, font=font_small)
            # oCy -> bit 1
            draw.line([(fa_x+fa_width, fa_y+80), (add_box_x, add_box_y+45)], fill=red, width=2)
            draw.text((fa_x+fa_width+40, fa_y+65), 'addC', fill=red, font=font_small)
            
            # subResult [1:0]
            draw.rectangle([sub_box_x, sub_box_y, sub_box_x+signal_box_w, sub_box_y+signal_box_h], 
                          outline=black, fill=light_gray, width=2)
            draw.text((sub_box_x+8, sub_box_y+5), '0', fill=black, font=font_medium)
            draw.text((sub_box_x+8, sub_box_y+30), '1', fill=black, font=font_medium)
            draw.text((sub_box_x-80, sub_box_y+20), 'subResult', fill=black, font=font_small)
            
            # FS出力 -> subResult
            # oD -> bit 0
            draw.line([(fs_x+fs_width, fs_y+50), (sub_box_x, sub_box_y+15)], fill=red, width=2)
            draw.text((fs_x+fs_width+40, fs_y+35), 'subD', fill=red, font=font_small)
            # oBr -> bit 1
            draw.line([(fs_x+fs_width, fs_y+80), (sub_box_x, sub_box_y+45)], fill=red, width=2)
            draw.text((fs_x+fs_width+40, fs_y+65), 'subD', fill=red, font=font_small)
            
            # === マルチプレクサ（三角形） ===
            mux_points = [
                (mux_x, mux_y),
                (mux_x, mux_y + mux_height),
                (mux_x + mux_width, mux_y + mux_height//2)
            ]
            draw.polygon(mux_points, outline=black, fill=(250, 250, 220), width=2)
            draw.text((mux_x+5, mux_y+mux_height//2-10), 'MUX', fill=black, font=font_medium)
            
            # addResult -> MUX (上側入力)
            draw.line([(add_box_x+signal_box_w, add_box_y+30), (mux_x, mux_y+40)], fill=red, width=2)
            draw.text((add_box_x+signal_box_w+10, add_box_y+15), '2', fill=black, font=font_small)
            
            # subResult -> MUX (下側入力)
            draw.line([(sub_box_x+signal_box_w, sub_box_y+30), (mux_x, mux_y+120)], fill=red, width=2)
            draw.text((sub_box_x+signal_box_w+10, sub_box_y+15), '2', fill=black, font=font_small)
            
            # SW4 -> MUX制御
            sw4_control_x = 700
            draw.line([(sw_start_x+sw_width, sw4_y+20), (sw4_control_x, sw4_y+20), 
                      (sw4_control_x, mux_y-30), (mux_x+20, mux_y-30), (mux_x+20, mux_y)], 
                     fill=red, width=2)
            draw.text((sw4_control_x-40, sw4_y+5), 'SW4', fill=red, font=font_small)
            
            # MUX出力 -> muxResult
            mux_out_x = mux_x + mux_width
            mux_out_y = mux_y + mux_height//2
            draw.line([(mux_out_x, mux_out_y), (not_x-100, mux_out_y)], fill=red, width=2)
            draw.text((mux_out_x+10, mux_out_y-20), 'muxResult', fill=black, font=font_small)
            draw.text((mux_out_x+10, mux_out_y+5), '2', fill=black, font=font_small)
            
            # === NOT gates (反転) ===
            # 上のNOT (bit 0)
            not1_center_x = not_x + not_size//2
            not1_center_y = not_y_top + not_size//2
            draw.polygon([(not_x, not_y_top), (not_x, not_y_top+not_size), 
                         (not_x+not_size, not1_center_y)], outline=black, fill=light_gray, width=2)
            draw.ellipse([not_x+not_size-2, not1_center_y-4, not_x+not_size+6, not1_center_y+4], 
                        outline=black, fill='white', width=2)
            draw.text((not_x+35, not_y_top+30), '0', fill=black, font=font_small)
            
            # 下のNOT (bit 1)
            not2_center_y = not_y_bot + not_size//2
            draw.polygon([(not_x, not_y_bot), (not_x, not_y_bot+not_size), 
                         (not_x+not_size, not2_center_y)], outline=black, fill=light_gray, width=2)
            draw.ellipse([not_x+not_size-2, not2_center_y-4, not_x+not_size+6, not2_center_y+4], 
                        outline=black, fill='white', width=2)
            draw.text((not_x+35, not_y_bot+30), '1', fill=black, font=font_small)
            
            # muxResult -> NOT gates
            split_x = not_x - 50
            draw.line([(not_x-100, mux_out_y), (split_x, mux_out_y)], fill=red, width=2)
            draw.ellipse([split_x-4, mux_out_y-4, split_x+4, mux_out_y+4], fill=red)
            draw.line([(split_x, mux_out_y), (split_x, not1_center_y), (not_x, not1_center_y)], fill=red, width=2)
            draw.line([(split_x, mux_out_y), (split_x, not2_center_y), (not_x, not2_center_y)], fill=red, width=2)
            
            # === 出力LED ===
            # LD1 (上)
            draw.ellipse([led_x, ld1_y, led_x+led_size, ld1_y+led_size], 
                        outline=black, fill=green, width=3)
            draw.text((led_x+led_size+10, ld1_y+15), 'LD1', fill=black, font=font_large)
            
            # LD2 (下)
            draw.ellipse([led_x, ld2_y, led_x+led_size, ld2_y+led_size], 
                        outline=black, fill=green, width=3)
            draw.text((led_x+led_size+10, ld2_y+15), 'LD2', fill=black, font=font_large)
            
            # NOT -> LED配線
            draw.line([(not_x+not_size+6, not1_center_y), (led_x, ld1_y+led_size//2)], fill=green, width=3)
            draw.line([(not_x+not_size+6, not2_center_y), (led_x, ld2_y+led_size//2)], fill=green, width=3)
            draw.text((not_x+not_size+40, not1_center_y-20), '2', fill=black, font=font_small)
            draw.text((not_x+not_size+40, not2_center_y+10), '2', fill=black, font=font_small)
            
            # 保存前に画像を確認用に一時保存
            output_path = self.img_dir / 'block_diagram_generated.png'
            img.save(str(output_path))
            print(f"ブロック図を生成しました: {output_path}")
            print("画像の確認を行っています...")
            
            # 画像の検証
            if self.validate_generated_image(output_path):
                print("✓ ブロック図の検証に成功しました")
                return output_path
            else:
                print("⚠ ブロック図に問題がある可能性があります")
                return output_path
            
        except ImportError:
            print("警告: Pillowライブラリがインストールされていないため、ブロック図を作成できません。")
            return None
        except Exception as e:
            print(f"ブロック図の作成中にエラーが発生しました: {e}")
            return None
    
    def validate_generated_image(self, image_path):
        """生成された画像の妥当性を確認"""
        try:
            from PIL import Image
            img = Image.open(image_path)
            
            # 基本的な検証
            width, height = img.size
            if width < 100 or height < 100:
                print("✗ 画像サイズが小さすぎます")
                return False
            
            # 画像が真っ白または真っ黒でないか確認
            pixels = list(img.getdata())
            unique_colors = len(set(pixels))
            if unique_colors < 10:
                print("✗ 画像の色数が少なすぎます（描画されていない可能性）")
                return False
            
            print(f"✓ 画像サイズ: {width}x{height}")
            print(f"✓ ユニークな色数: {unique_colors}")
            
            return True
            
        except Exception as e:
            print(f"✗ 画像検証中にエラー: {e}")
            return False
        
    def generate(self):
        """レポート生成のメイン処理"""
        print("レポート生成を開始します...")
        
        # ヘッダー情報
        self.add_header_info()
        
        # 各章の追加
        print("第1章: 目的")
        self.add_chapter1_purpose()
        
        print("第2章: 原理・設計")
        self.add_chapter2_principle()
        
        print("第3章: 実装")
        self.add_chapter3_implementation()
        
        print("第4章: シミュレーション結果")
        self.add_chapter4_simulation()
        
        print("第5章: 考察")
        self.add_chapter5_discussion()
        
        # ファイル保存
        output_path = self.base_dir / self.output_filename
        self.doc.save(str(output_path))
        print(f"\nレポートを生成しました: {output_path}")
        
        return output_path


def main():
    """メイン関数"""
    generator = ReportGenerator()
    generator.generate()


if __name__ == "__main__":
    main()
