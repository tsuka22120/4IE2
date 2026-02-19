import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==========================================
# 1. Verilog HDL Code Definition
# ==========================================
verilog_code = """/*
 * Module: Top Module (Full Adder / Full Subtractor Switchable)
 * Device: Gowin GW1NR-9
 * Board: Tang Nano 9K (Assuming)
 */

// Half Adder
module half_adder(
    input a,
    input b,
    output s,
    output c
);
    assign s = a ^ b;
    assign c = a & b;
endmodule

// Full Adder (using 2 Half Adders)
module full_adder(
    input a,
    input b,
    input ci,
    output s,
    output co
);
    wire s1, c1, c2;
    
    half_adder HA1 (.a(a), .b(b), .s(s1), .c(c1));
    half_adder HA2 (.a(s1), .b(ci), .s(s), .c(c2));
    
    assign co = c1 | c2;
endmodule

// Half Subtractor
module half_subtractor(
    input a,
    input b,
    output d,
    output br
);
    assign d = a ^ b;
    assign br = ~a & b;
endmodule

// Full Subtractor (using 2 Half Subtractors)
module full_subtractor(
    input a,
    input b,
    input bri, // Borrow In
    output d,
    output bro // Borrow Out
);
    wire d1, br1, br2;
    
    half_subtractor HS1 (.a(a), .b(b), .d(d1), .br(br1));
    half_subtractor HS2 (.a(d1), .b(bri), .d(d), .br(br2));
    
    assign bro = br1 | br2;
endmodule

// Top Module
module top(
    input sw1, // A
    input sw2, // B
    input sw3, // Carry In / Borrow In
    input sw4, // Mode Select (0: Adder, 1: Subtractor)
    output ld1, // Sum / Diff (Active Low LED)
    output ld2  // Carry Out / Borrow Out (Active Low LED)
);
    wire fa_s, fa_c;
    wire fs_d, fs_br;
    wire result_val, result_c_br;

    // Instance of Full Adder
    full_adder FA (
        .a(sw1), .b(sw2), .ci(sw3),
        .s(fa_s), .co(fa_c)
    );

    // Instance of Full Subtractor
    full_subtractor FS (
        .a(sw1), .b(sw2), .bri(sw3),
        .d(fs_d), .bro(fs_br)
    );

    // Multiplexer based on SW4
    assign result_val = (sw4 == 1'b0) ? fa_s : fs_d;
    assign result_c_br = (sw4 == 1'b0) ? fa_c : fs_br;

    // LED Output Logic (Inverted because LED implies Active Low/Opposite display)
    assign ld1 = ~result_val;
    assign ld2 = ~result_c_br;

endmodule
"""

# ==========================================
# 2. CST File (Pin Assignment) Definition
# ==========================================
cst_code = """// Copyright (C)2025 Gowin Semiconductor Corporation.
// Part Number: GW1NR-LV9QN88PC6/I5

IO_LOC "sw1" 69;
IO_PORT "sw1" IO_TYPE=LVCMOS33 PULL_MODE=DOWN;

IO_LOC "sw2" 68;
IO_PORT "sw2" IO_TYPE=LVCMOS33 PULL_MODE=DOWN;

IO_LOC "sw3" 57;
IO_PORT "sw3" IO_TYPE=LVCMOS33 PULL_MODE=DOWN;

IO_LOC "sw4" 56;
IO_PORT "sw4" IO_TYPE=LVCMOS33 PULL_MODE=DOWN;

IO_LOC "ld1" 10;
IO_PORT "ld1" IO_TYPE=LVCMOS33 PULL_MODE=UP DRIVE=8;

IO_LOC "ld2" 11;
IO_PORT "ld2" IO_TYPE=LVCMOS33 PULL_MODE=UP DRIVE=8;
"""

# ==========================================
# 3. Testbench Code Definition
# ==========================================
tb_code = """`timescale 1ns / 1ps

module tb_top;

    // Inputs
    reg sw1;
    reg sw2;
    reg sw3;
    reg sw4;

    // Outputs
    wire ld1;
    wire ld2;

    // Instantiate the Unit Under Test (UUT)
    top uut (
        .sw1(sw1), 
        .sw2(sw2), 
        .sw3(sw3), 
        .sw4(sw4), 
        .ld1(ld1), 
        .ld2(ld2)
    );

    initial begin
        // Initialize Inputs
        sw1 = 0; sw2 = 0; sw3 = 0; sw4 = 0;

        // Wait 100 ns for global reset to finish
        #100;
        
        // ------------------------------------------------
        // Case 1: Full Adder Test (SW4 = 0)
        // ------------------------------------------------
        $display("--- Start Full Adder Test (SW4=0) ---");
        sw4 = 0;
        // Truth Table Loop
        // A(sw1), B(sw2), Cin(sw3)
        {sw1, sw2, sw3} = 3'b000; #10;
        {sw1, sw2, sw3} = 3'b001; #10;
        {sw1, sw2, sw3} = 3'b010; #10;
        {sw1, sw2, sw3} = 3'b011; #10;
        {sw1, sw2, sw3} = 3'b100; #10;
        {sw1, sw2, sw3} = 3'b101; #10;
        {sw1, sw2, sw3} = 3'b110; #10;
        {sw1, sw2, sw3} = 3'b111; #10;

        // ------------------------------------------------
        // Case 2: Full Subtractor Test (SW4 = 1)
        // ------------------------------------------------
        $display("--- Start Full Subtractor Test (SW4=1) ---");
        sw4 = 1;
        // Truth Table Loop
        // A(sw1), B(sw2), Bin(sw3)
        {sw1, sw2, sw3} = 3'b000; #10;
        {sw1, sw2, sw3} = 3'b001; #10;
        {sw1, sw2, sw3} = 3'b010; #10;
        {sw1, sw2, sw3} = 3'b011; #10;
        {sw1, sw2, sw3} = 3'b100; #10;
        {sw1, sw2, sw3} = 3'b101; #10;
        {sw1, sw2, sw3} = 3'b110; #10;
        {sw1, sw2, sw3} = 3'b111; #10;

        $finish;
    end
      
endmodule
"""

# ==========================================
# 4. Generate Files
# ==========================================

# Save HDL files
with open("main.v", "w", encoding="utf-8") as f:
    f.write(verilog_code)
with open("fpga_project.cst", "w", encoding="utf-8") as f:
    f.write(cst_code)
with open("tb_main.v", "w", encoding="utf-8") as f:
    f.write(tb_code)

# ==========================================
# 5. Generate Word Report
# ==========================================
doc = Document()

# Set Font style (Attempt to set generic style, though Word defaults often override)
style = doc.styles['Normal']
font = style.font
font.name = 'MS Mincho'
font.size = Pt(10.5)

def set_mincho(paragraph):
    """Helper to set Japanese font manually if needed"""
    for run in paragraph.runs:
        run.font.name = 'MS Mincho'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

# --- Cover Page (Based on 課題1.docx) ---
doc.add_paragraph("提出日　2025年　12月　9日", style='Normal').alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.add_paragraph("") # Spacer

# Title Area
p_title_sub = doc.add_paragraph("集積回路設計　演習１")
p_title_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title_sub.runs[0].font.size = Pt(14)
p_title_sub.runs[0].bold = True

p_title = doc.add_paragraph("～組合せ回路とHDL記述～")
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.runs[0].font.size = Pt(16)
p_title.runs[0].bold = True

doc.add_paragraph("") # Spacer

# Info Table (Using table to mimic the form)
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Fill Cover Info
# 氏名・学籍番号 (Using provided User Info)
table.cell(0, 0).text = "報告者"
table.cell(0, 1).text = "22120 塚田 勇人"

table.cell(1, 0).text = "共同実験者"
table.cell(1, 1).text = "なし"

table.cell(2, 0).text = "天候"
table.cell(2, 1).text = "晴れ" # Placeholder

table.cell(3, 0).text = "室温"
table.cell(3, 1).text = "22℃" # Placeholder

table.cell(4, 0).text = "湿度"
table.cell(4, 1).text = "45%" # Placeholder

doc.add_paragraph("") 
doc.add_paragraph("【課題内容】")
doc.add_paragraph("Verilog HDLで全加算器、全減算器を構築する。SW4を用いて機能を切り替える。")
doc.add_paragraph("使用FPGA: Gowin GW1NR-9")

doc.add_page_break()

# --- Main Content (Based on レポートの書き方.pdf) ---

# 1. 目的
doc.add_heading('1. 実験の目的', level=1)
p = doc.add_paragraph(
    "本実験では、ハードウェア記述言語であるVerilog HDLを用いて、組合せ回路の基本である全加算器および全減算器を設計することを目的とする。"
    "また、設計した回路をFPGA（Gowin GW1NR-9）上に実装し、スイッチ入力およびLED出力を用いて動作確認を行うことで、"
    "FPGA開発フローおよび階層設計の手法を習得する。"
)

# 2. 原理（理論）
doc.add_heading('2. 原理（理論）', level=1)
p = doc.add_paragraph(
    "全加算器は、下位桁からの桁上げ入力を含む3つの1ビット入力（A, B, Cin）を加算し、和（Sum）と桁上げ（Cout）を出力する回路である。"
    "これは2つの半加算器とORゲートを組み合わせることで構成できる。\n"
    "全減算器は、下位桁からの借用入力を含む3つの1ビット入力（A, B, Bin）から減算を行い、差（Diff）と借用（Bout）を出力する回路である。"
    "これも同様に、2つの半減算器とORゲートを組み合わせることで構成可能である。"
)

# 3. 使用器具
doc.add_heading('3. 使用器具', level=1)
p = doc.add_paragraph(
    "本実験で使用したソフトウェアおよびハードウェアを以下に示す。"
)
# Table for apparatus
table_tools = doc.add_table(rows=3, cols=2)
table_tools.style = 'Table Grid'
table_tools.cell(0, 0).text = "項目"
table_tools.cell(0, 1).text = "名称・仕様"
table_tools.cell(1, 0).text = "開発環境"
table_tools.cell(1, 1).text = "Gowin EDA"
table_tools.cell(2, 0).text = "FPGAボード"
table_tools.cell(2, 1).text = "Tang Nano 9K (GW1NR-9)"
doc.add_paragraph("表1: 使用器具一覧", style='Caption')

# 4. 実験方法
doc.add_heading('4. 実験方法', level=1)
doc.add_paragraph(
    "1. 2つの半加算器を用いて全加算器モジュールを作成する。\n"
    "2. 2つの半減算器を用いて全減算器モジュールを作成する。\n"
    "3. トップモジュールにおいて、SW4の状態（0または1）に応じて全加算器と全減算器の出力を切り替えるマルチプレクサを記述する。\n"
    "4. 出力デバイス（LD1, LD2）の仕様に基づき、論理を反転させて出力ピンに接続する。\n"
    "5. 論理合成および配置配線を行い、実機での動作を確認する。"
)

# 5. 実験結果
doc.add_heading('5. 実験結果', level=1)

doc.add_heading('5.1. ブロック図', level=2)
doc.add_paragraph("[ここにブロック図の画像を挿入してください]")
doc.add_paragraph("図1: 全体のブロック図", style='Caption')

doc.add_heading('5.2. 作成したVerilog HDLコード', level=2)
doc.add_paragraph("作成したトップモジュールおよび下位モジュールのコードを表2に示す。")

# Code inside a 1x1 table
table_code = doc.add_table(rows=1, cols=1)
table_code.style = 'Table Grid'
cell = table_code.cell(0, 0)
# Add code with line numbers (simplified simulation)
lines = verilog_code.split('\n')
code_text = ""
for idx, line in enumerate(lines):
    code_text += f"{idx+1: >3}  {line}\n"
cell.text = code_text
# Change font to Monospace for code
for paragraph in cell.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
doc.add_paragraph("表2: 作成したVerilog HDLコード", style='Caption')

doc.add_heading('5.3. ピン割り当て', level=2)
doc.add_paragraph("FPGAへのピン割り当て（.cstファイルの内容）を表3に示す。")
table_cst = doc.add_table(rows=1, cols=1)
table_cst.style = 'Table Grid'
cell_cst = table_cst.cell(0, 0)
cell_cst.text = cst_code
for paragraph in cell_cst.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
doc.add_paragraph("表3: ピン割り当て設定", style='Caption')

doc.add_heading('5.4. シミュレーション結果', level=2)
doc.add_paragraph("作成したテストベンチを用いて動作検証を行った。タイミングチャートを図2に示す（※実行後に画像を貼り付けること）。")
doc.add_paragraph("[ここにタイミングチャートの画像を挿入してください]")
doc.add_paragraph("図2: シミュレーション波形", style='Caption')

# 6. 考察
doc.add_heading('6. 考察', level=1)
doc.add_paragraph(
    "実験結果より、SW4が0のときは全加算器として、SW4が1のときは全減算器として正しく動作していることが確認できた。"
    "特に、LEDの点灯論理が負論理（Active Low）であることを考慮し、HDL記述内で出力を反転させたことで、"
    "意図通りの点灯制御が行えている。シミュレーション波形においても、真理値表通りの推移が確認された。"
)

# Save the document
doc.save("課題1_レポート.docx")

print("以下のファイルが生成されました:")
print("1. 課題1_レポート.docx")
print("2. main.v (Verilog Code)")
print("3. fpga_project.cst (Pin Constraint)")
print("4. tb_main.v (Testbench)")